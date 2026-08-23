"""Uploaded resumes must not carry their authoring metadata into our storage.

These tests build real PDF and DOCX files with metadata set, run them through the
redaction path, and assert two things that have to hold together:

1. the identifying metadata is gone, and
2. the document text still extracts.

The second assertion matters as much as the first. Redaction rewrites the file, and
a rewrite that quietly destroys the text would disable resume parsing and the
resume-readiness review while every upload still returned 200.
"""
from __future__ import annotations

import re

from io import BytesIO

import pytest

from app.services.document_redaction import strip_document_metadata

pypdf = pytest.importorskip("pypdf", reason="pypdf is declared in requirements.txt")
docx = pytest.importorskip("docx", reason="python-docx is declared in requirements.txt")


def _build_pdf_with_metadata() -> bytes:
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    writer.add_metadata(
        {
            "/Author": "Priya Sharma",
            "/Creator": "Acme Corp Internal Template",
            "/Producer": "Confidential Employer Word Build",
            "/Title": "Priya CV final v3",
        }
    )
    buffer = BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


def _build_docx_with_metadata(text: str = "Python SQL AWS internship") -> bytes:
    from docx import Document

    document = Document()
    document.add_paragraph(text)
    properties = document.core_properties
    properties.author = "Priya Sharma"
    properties.last_modified_by = "acme-laptop-07"
    properties.company = "Acme Corp" if hasattr(properties, "company") else None
    properties.title = "Priya CV final v3"
    properties.comments = "reviewed by manager"
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


class TestPdfRedaction:
    def test_author_and_producer_are_removed(self):
        original = _build_pdf_with_metadata()
        from pypdf import PdfReader

        before = PdfReader(BytesIO(original)).metadata or {}
        assert before.get("/Author") == "Priya Sharma", "fixture did not set metadata"

        redacted, was_redacted = strip_document_metadata(extension=".pdf", content=original)
        assert was_redacted is True

        after = PdfReader(BytesIO(redacted)).metadata or {}
        values = " ".join(str(value) for value in after.values())
        assert "Priya Sharma" not in values
        assert "Acme Corp" not in values
        assert "Confidential Employer" not in values

    def test_page_count_survives_redaction(self):
        from pypdf import PdfReader

        original = _build_pdf_with_metadata()
        redacted, _ = strip_document_metadata(extension=".pdf", content=original)
        assert len(PdfReader(BytesIO(redacted)).pages) == len(PdfReader(BytesIO(original)).pages)

    def test_malformed_pdf_falls_open_to_the_original(self):
        """A student who cannot upload a CV is worse off than one whose metadata stayed."""
        junk = b"%PDF-1.4 this is not really a pdf"
        result, was_redacted = strip_document_metadata(extension=".pdf", content=junk)
        assert result == junk
        assert was_redacted is False


class TestDocxRedaction:
    def test_author_and_editor_are_removed(self):
        from docx import Document

        original = _build_docx_with_metadata()
        assert Document(BytesIO(original)).core_properties.author == "Priya Sharma"

        redacted, was_redacted = strip_document_metadata(extension=".docx", content=original)
        assert was_redacted is True

        properties = Document(BytesIO(redacted)).core_properties
        assert properties.author in ("", None)
        assert properties.last_modified_by in ("", None)
        assert properties.title in ("", None)
        assert properties.comments in ("", None)

    def test_document_text_survives_redaction(self):
        """The reason this test exists: a rewrite that eats the text breaks parsing."""
        from docx import Document

        original = _build_docx_with_metadata("Python SQL AWS internship")
        redacted, _ = strip_document_metadata(extension=".docx", content=original)
        paragraphs = [p.text for p in Document(BytesIO(redacted)).paragraphs if p.text]
        assert "Python SQL AWS internship" in paragraphs

    def test_malformed_docx_falls_open_to_the_original(self):
        junk = b"PK\x03\x04 not really a docx"
        result, was_redacted = strip_document_metadata(extension=".docx", content=junk)
        assert result == junk
        assert was_redacted is False


class TestPassthroughFormats:
    def test_txt_is_returned_untouched(self):
        content = b"Priya Sharma\nPython, SQL\n"
        result, was_redacted = strip_document_metadata(extension=".txt", content=content)
        assert result == content
        assert was_redacted is False

    def test_legacy_doc_is_passed_through_not_corrupted(self):
        """Legacy .doc has no parser in our dependency set; it must not be mangled."""
        content = b"\xd0\xcf\x11\xe0 legacy ole stream"
        result, was_redacted = strip_document_metadata(extension=".doc", content=content)
        assert result == content
        assert was_redacted is False


class TestUploadPathUsesRedaction:
    def test_upload_endpoint_writes_the_redacted_bytes(self):
        """Guards against the redaction call being added and then bypassed."""
        import inspect

        from app.api.api_v1.endpoints.users import upload_resume

        source = inspect.getsource(upload_resume)
        assert "strip_document_metadata" in source

        # Match the argument, not the call syntax. This asserted the exact string
        # "storage_path.write_bytes(storable)" and so failed the moment the write
        # was moved off the event loop with asyncio.to_thread - a change that did
        # not touch which bytes get written. What actually matters is that the
        # redacted `storable` reaches disk and the raw `content` never does.
        write_calls = re.findall(r"storage_path\.write_bytes[,(]\s*([A-Za-z_][A-Za-z0-9_]*)", source)
        assert write_calls, "upload_resume must write the resume to storage_path."
        assert all(arg == "storable" for arg in write_calls), (
            "upload_resume must persist the redacted bytes, not the uploaded ones. "
            f"write_bytes received: {write_calls}"
        )
